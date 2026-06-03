"""Generate zoology-themed sample files matching the assignment spec.

Creates 8 files under data/samples/zoology/:
  - 4 PDFs  (.pdf)
  - 1 DOCX  (.docx)
  - 1 XLS   (.xls)
  - 1 TXT   (.txt)
  - 1 JPG   (.jpg)

All content is non-technical and themed around zoology/biology.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from xlwt import Workbook


def write_samples(root: Path) -> None:
    root.mkdir(parents=True, exist_ok=True)

    # ── 4 PDFs ────────────────────────────────────────────────────────────────

    pdf_specs = [
        (
            "african_elephant_study.pdf",
            "African Elephant Population Dynamics",
            (
                "Loxodonta africana herds across the Amboseli basin were studied "
                "from 2018 to 2023. Population counts show a 12% recovery trend "
                "following anti-poaching measures. Key findings: herd sizes increased "
                "from 42 to 67 individuals in the northern sector."
            ),
        ),
        (
            "marine_mammals_report.pdf",
            "Cetacean Survey: Pacific Northwest Corridor",
            (
                "Dolphins and orca pods were monitored via hydrophone arrays along "
                "the Pacific Northwest corridor. Resident J-Pod was documented with "
                "24 individuals. Echolocation click trains recorded at depths of 200 m."
            ),
        ),
        (
            "bird_migration_analysis.pdf",
            "Migratory Patterns of Arctic Terns",
            (
                "Sterna paradisaea individuals were tracked via geolocators across "
                "their 70,000 km annual migration route from Arctic breeding grounds "
                "to Antarctic wintering sites. Stopover duration averaged 8.3 days."
            ),
        ),
        (
            "amphibian_survey_2023.pdf",
            "Global Amphibian Decline: Chytrid Fungus Impact",
            (
                "Batrachochytrium dendrobatidis was detected in 501 species across "
                "54 countries in this 2023 survey. Mortality rates exceeded 80% in "
                "affected populations of Atelopus species in Central America."
            ),
        ),
    ]

    for filename, title, body in pdf_specs:
        pdf_canvas = canvas.Canvas(str(root / filename), pagesize=letter)
        pdf_canvas.setTitle(title)
        pdf_canvas.setFont("Helvetica-Bold", 16)
        pdf_canvas.drawString(72, 750, title)
        pdf_canvas.setFont("Helvetica", 11)
        # word-wrap the body text manually at ~80 chars
        words = body.split()
        line, lines = [], []
        for w in words:
            if len(" ".join(line + [w])) <= 80:
                line.append(w)
            else:
                lines.append(" ".join(line))
                line = [w]
        if line:
            lines.append(" ".join(line))
        y = 720
        for ln in lines:
            pdf_canvas.drawString(72, y, ln)
            y -= 16
            if y < 72:
                pdf_canvas.showPage()
                pdf_canvas.setFont("Helvetica", 11)
                y = 750
        pdf_canvas.save()

    # ── DOCX: coral_reef_observations.docx ─────────────────────────────────
    doc = Document()
    doc.add_heading("Coral Reef Biodiversity Observations", level=1)
    doc.add_paragraph(
        "Transect surveys were conducted at Great Barrier Reef Site 7B during the "
        "2023 field season. A total of 285 coral species were identified across 12 "
        "genera. Bleaching events were recorded in Sectors 3 and 5 following the "
        "February thermal anomaly. Recommendations include targeted monitoring of "
        "Acropora populations in the southern sector."
    )
    doc.save(str(root / "coral_reef_observations.docx"))

    # ── XLS: species_count_2024.xls (xlwt produces Excel 97-2003 binary) ─────
    book = Workbook()
    sheet = book.add_sheet("Species Count")
    headers = ["Species", "Region", "Count_2022", "Count_2023", "IUCN Status"]
    data = [
        ("Snow Leopard", "Himalayas", 4500, 4700, "Vulnerable"),
        ("Giant Panda", "Sichuan Basin", 1864, 1900, "Vulnerable"),
        ("Amur Tiger", "Far East Russia", 500, 537, "Endangered"),
        ("Blue Whale", "Pacific", 10000, 10200, "Endangered"),
        ("Hawksbill Turtle", "Caribbean", 8200, 8300, "Critically Endangered"),
    ]
    for col, h in enumerate(headers):
        sheet.write(0, col, h)
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, val in enumerate(row_data):
            sheet.write(row_idx, col_idx, val)
    xls_path = root / "species_count_2024.xls"
    book.save(str(xls_path))
    # Remove legacy misnamed .xlsx if present (xlwt binary saved as .xlsx will not open)
    legacy_xlsx = root / "species_count_2024.xlsx"
    if legacy_xlsx.exists():
        legacy_xlsx.unlink()

    # ── TXT: field_notes_borneo.txt ────────────────────────────────────────
    (root / "field_notes_borneo.txt").write_text(
        "Field Notes - Borneo Rainforest Expedition 2024\n"
        "Date: March 12-28, 2024\n"
        "Observer: Dr. S. Patel, Wildlife Conservation Society\n\n"
        "Day 1: Proboscis monkey troop (11 individuals) observed at riverside foraging zone.\n"
        "Day 3: Pygmy elephant tracks near salt lick at GPS 5.2103N, 117.4432E.\n"
        "Day 7: Orangutan nest count: 34 active nests recorded in 2 km transect.\n"
        "Day 12: Sun bear claw marks identified on Dipterocarp tree at elevation 340 m.\n"
        "Day 14: Faecal samples from 8 sambar deer collected for dietary analysis.\n",
        encoding="utf-8",
    )

    # ── JPG: jaguar_photo_rainforest.jpg (camera-trap style) ───────────────
    img = Image.new("RGB", (800, 600), color=(34, 90, 34))  # rainforest green
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 800, 80], fill=(20, 50, 20))
    draw.text(
        (20, 20),
        "Jaguar - Panthera onca",
        fill=(220, 220, 150),
    )
    draw.text(
        (20, 45),
        "Rainforest Camera Trap | Brazil | 2024",
        fill=(180, 180, 120),
    )
    # Simple silhouette suggestion
    draw.ellipse([300, 200, 500, 450], fill=(160, 120, 60))
    draw.ellipse([340, 150, 460, 220], fill=(160, 120, 60))  # head
    img.save(str(root / "jaguar_photo_rainforest.jpg"), format="JPEG", quality=85)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        target = Path(sys.argv[1])
    else:
        target = Path(__file__).resolve().parents[1] / "data" / "samples" / "zoology"
    write_samples(target)
    print(f"Sample data written to {target}")
